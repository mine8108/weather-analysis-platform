"""
报告导出模块：图表 PNG 导出、增强 Word 分析报告、数据 CSV 导出

设计要点：
- 专业版：完整章节 + 表格 + 图表 + 防御建议，专业受众可作为正式交付物
- 通俗版：叙述式 + 生活语言 + 问答式，适合非专业用户
- fc_df 图表支持：修复「仅有 GFS 预报时报告无图」的根因
- 排版规范：封面表 + 目录 + 子标题 + 表格 + 分页 + 配色
"""

from io import BytesIO
from datetime import datetime
import traceback

import pandas as pd
import streamlit as st
import plotly.graph_objects as go

# docx 相关依赖（顶层 import 失败时不影响 streamlit 加载）
try:
    from docx.shared import RGBColor
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    RGBColor = Pt = Cm = Inches = WD_ALIGN_PARAGRAPH = qn = OxmlElement = None
    _DOCX_OK = False

# 报告用：字段名别名映射（预报字段 → 标准字段）
_FIELD_ALIAS_REPORT = {
    "temperature_2m": "temperature",
    "relative_humidity_2m": "humidity",
    "surface_pressure": "pressure",
    "wind_speed_10m": "wind_speed",
    "precipitation_sum": "precipitation",
    "precipitation_hours": "precipitation",
}

# 报告排版常量
_TITLE_COLOR = "1E3A5F"           # 标题深蓝
_HEADER_BG = "DBEAFE"             # 表头底色浅蓝
_HEADER_FG = "1E3A5F"             # 表头字色深蓝
_SUBTLE_COLOR = "6B7280"          # 副文本灰


# ============================================================
# 排版辅助函数
# ============================================================
def _set_default_styles(doc):
    """设置文档默认样式：正文字号、字体、段间距"""
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3


def _add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def _build_cover_table(doc, rows, col_widths=(4, 11)):
    """封面信息表：rows=[(label, value), ...]"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)

    for i, (label, value) in enumerate(rows):
        # 左列标签
        cell_l = table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(str(label))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(_HEADER_FG)
        # 右列值
        cell_r = table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(str(value))
        run.font.size = Pt(11)
    return table


def _add_table(doc, headers, rows, col_widths=None, header_bg=True):
    """通用表格渲染：headers=列名列表，rows=[[cell, ...], ...]"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(_HEADER_FG)
        if header_bg:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), _HEADER_BG)
            tc_pr.append(shd)

    # 数据行
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def _build_toc(doc, sections):
    """目录：sections=[(title, page_hint), ...]"""
    doc.add_heading("目录", level=1)
    for title, _page in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"  •  {title}")
        run.font.size = Pt(12)
    doc.add_paragraph("")


# ============================================================
# 字段别名解析
# ============================================================
def _resolve_field(df, field):
    """按别名解析字段，返回实际列名或 None"""
    if field in df.columns:
        return field
    for alias, target in _FIELD_ALIAS_REPORT.items():
        if target == field and alias in df.columns:
            return alias
    return None


def _resolve_stats(df, stats_fields):
    """解析统计字段（支持别名），返回 {field: Series}"""
    result = {}
    for f in stats_fields:
        actual = _resolve_field(df, f)
        if actual and not df[actual].dropna().empty:
            result[f] = df[actual].dropna()
    return result


# ============================================================
# 工具函数
# ============================================================
def _number(n):
    """中文数字映射"""
    nums = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
    return nums[n] if n < len(nums) else str(n)


def _to_image_safe(fig, timeout=15, **kwargs):
    """通过子进程超时保护调用 fig.to_image()。
    kaleido 0.2.1 在缺 chromium 环境下会挂死整个 Python 进程，
    用 multiprocessing 在子进程中执行 + 超时杀死来防止整个应用崩溃。
    """
    import multiprocessing
    import pickle

    # plotly Figure 不能直接 pickle，改用 to_dict 后再传
    fig_dict = fig.to_dict()

    def _worker(q, fd, kw):
        try:
            import plotly.graph_objects as go
            f = go.Figure(fd)
            png = f.to_image(**kw)
            q.put(("ok", png))
        except Exception as e:
            q.put(("err", str(e)))

    ctx = multiprocessing.get_context("spawn")
    q = ctx.Queue()
    p = ctx.Process(target=_worker, args=(q, fig_dict, kwargs))
    p.start()
    p.join(timeout)
    if p.is_alive():
        p.terminate()
        p.join(2)
        if p.is_alive():
            p.kill()
            p.join()
        return None
    try:
        status, data = q.get_nowait()
        return data if status == "ok" else None
    except Exception:
        return None


def export_chart_as_png(fig, filename="chart.png"):
    """导出 Plotly 图为 PNG 字节流(带超时保护)"""
    if fig is None:
        return None
    try:
        return _to_image_safe(fig, timeout=15,
                              format="png", scale=2, width=1200, height=800)
    except Exception:
        return None


def export_data_csv(df):
    """导出数据为 CSV 字节流"""
    if df is None or df.empty:
        return None
    return df.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")


def _append_data_support(doc, warn):
    """在防御建议条文后追加『本次数据支撑』附注，消除纯套话感。"""
    detail = warn.get("detail", "")
    if not detail:
        return
    try:
        p = doc.add_paragraph(f"    数据支撑：{detail}")
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(120, 120, 120)
    except Exception:
        pass


# ============================================================
# 图表生成
# ============================================================
def _hex_to_rgba(hex_color, alpha=0.2):
    """将 #RRGGBB 转为 rgba(r, g, b, a) 字符串"""
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return hex_color
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    return f"rgba({r}, {g}, {b}, {alpha})"


def _build_one_chart(x, y, name, title, color, kind="line", fill=False):
    """生成单张图表的辅助函数。kind: line / bar
    注意：不要设置 font.family，不要用 8 位 hex fillcolor——kaleido 0.2.1 + 新版 plotly 会崩溃挂死。
    """
    if kind == "bar":
        fig = go.Figure(go.Bar(x=x, y=y, marker_color=color, name=name))
    else:
        fill_mode = "tozeroy" if fill else None
        fig = go.Figure(go.Scatter(
            x=x, y=y, mode="lines", name=name,
            line=dict(color=color, width=2),
            fill=fill_mode, fillcolor=_hex_to_rgba(color, 0.2) if fill else None,
        ))
    fig.update_layout(
        title=title, height=320,
        margin=dict(l=40, r=20, t=40, b=40),
        plot_bgcolor="white", paper_bgcolor="white",
        xaxis=dict(gridcolor="#e5e7eb"),
        yaxis=dict(gridcolor="#e5e7eb"),
    )
    return fig


def _generate_report_charts(df=None, forecast_df=None):
    """从观测数据和/或 GFS 预报生成图表。
    返回 dict,key 格式为 `前缀:图表名`,前缀 `obs` 表观测,`fc` 表预报。
    """
    figs = {}

    # ---- 观测数据图表 ----
    if df is not None and not df.empty:
        try:
            x = df["timestamp"] if "timestamp" in df.columns else df.index

            temp_col = _resolve_field(df, "temperature")
            if temp_col and df[temp_col].dropna().size > 0:
                figs["obs:temperature"] = _build_one_chart(
                    x, df[temp_col], "气温", "观测气温时序", "#e74c3c")

            prec_col = _resolve_field(df, "precipitation")
            if prec_col and df[prec_col].dropna().size > 0:
                if "timestamp" in df.columns:
                    daily = df.copy()
                    daily["date"] = daily["timestamp"].dt.date
                    daily_p = daily.groupby("date")[prec_col].sum()
                    x_d = [str(d) for d in daily_p.index]
                else:
                    x_d, daily_p = list(range(len(df))), df[prec_col]
                figs["obs:precipitation"] = _build_one_chart(
                    x_d, daily_p.values, "降水", "观测逐日降水量", "#2980b9",
                    kind="bar")

            pres_col = _resolve_field(df, "pressure")
            if pres_col and df[pres_col].dropna().size > 0:
                figs["obs:pressure"] = _build_one_chart(
                    x, df[pres_col], "气压", "观测气压时序", "#27ae60")

            wind_col = _resolve_field(df, "wind_speed")
            if wind_col and df[wind_col].dropna().size > 0:
                figs["obs:wind_speed"] = _build_one_chart(
                    x, df[wind_col], "风速", "观测风速时序", "#8e44ad")
        except Exception:
            pass

    # ---- GFS 预报图表 ----
    if forecast_df is not None and not forecast_df.empty:
        try:
            x = (forecast_df["timestamp"]
                 if "timestamp" in forecast_df.columns
                 else forecast_df.index)

            # 1. 气温 + 体感温度叠加
            temp_col = _resolve_field(forecast_df, "temperature")
            if temp_col and forecast_df[temp_col].dropna().size > 0:
                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    x=x, y=forecast_df[temp_col], mode="lines", name="气温",
                    line=dict(color="#e74c3c", width=2)))
                app_col = _resolve_field(forecast_df, "apparent_temperature")
                if app_col and forecast_df[app_col].dropna().size > 0:
                    fig.add_trace(go.Scatter(
                        x=x, y=forecast_df[app_col], mode="lines",
                        name="体感温度",
                        line=dict(color="#f97316", width=2, dash="dash")))
                fig.update_layout(
                    title="GFS 预报气温 / 体感温度", height=320,
                    margin=dict(l=40, r=20, t=40, b=40),
                    plot_bgcolor="white", paper_bgcolor="white",
                    xaxis=dict(gridcolor="#e5e7eb"),
                    yaxis=dict(gridcolor="#e5e7eb"),
                )
                figs["fc:temperature"] = fig

            # 2. 降水(逐时柱)
            prec_col = _resolve_field(forecast_df, "precipitation")
            if prec_col and forecast_df[prec_col].dropna().size > 0:
                fig = go.Figure(go.Bar(
                    x=x, y=forecast_df[prec_col],
                    marker_color="#2980b9", name="逐时降水"))
                fig.update_layout(
                    title="GFS 预报逐时降水量", height=320,
                    margin=dict(l=40, r=20, t=40, b=40),
                    plot_bgcolor="white", paper_bgcolor="white",
                    xaxis=dict(gridcolor="#e5e7eb"),
                    yaxis=dict(gridcolor="#e5e7eb"),
                )
                figs["fc:precipitation"] = fig

            # 3. 风速时序
            wind_col = _resolve_field(forecast_df, "wind_speed")
            if wind_col and forecast_df[wind_col].dropna().size > 0:
                figs["fc:wind_speed"] = _build_one_chart(
                    x, forecast_df[wind_col], "风速",
                    "GFS 预报风速时序", "#8e44ad")

            # 4. 相对湿度时序
            hum_col = _resolve_field(forecast_df, "humidity")
            if hum_col and forecast_df[hum_col].dropna().size > 0:
                figs["fc:humidity"] = _build_one_chart(
                    x, forecast_df[hum_col], "湿度",
                    "GFS 预报相对湿度", "#06b6d4", fill=True)

            # 5. 降水概率时序
            prob_col = _resolve_field(forecast_df, "precipitation_probability")
            if prob_col and forecast_df[prob_col].dropna().size > 0:
                figs["fc:precip_prob"] = _build_one_chart(
                    x, forecast_df[prob_col], "降水概率",
                    "GFS 预报降水概率(%)", "#0ea5e9", fill=True)
        except Exception:
            pass

    return figs


# ============================================================
# 图表标题映射
# ============================================================
_CHART_CAPTIONS = {
    "obs:temperature": "观测气温时序变化",
    "obs:precipitation": "观测逐日降水量分布",
    "obs:pressure": "观测气压时序变化",
    "obs:wind_speed": "观测风速时序变化",
    "fc:temperature": "GFS 预报气温 / 体感温度",
    "fc:precipitation": "GFS 预报逐时降水量",
    "fc:wind_speed": "GFS 预报风速时序",
    "fc:humidity": "GFS 预报相对湿度时序",
    "fc:precip_prob": "GFS 预报降水概率",
}


def _insert_chart(doc, key, fig):
    """插入单张图表：图片 + 居中图题，失败返回错误信息"""
    caption = _CHART_CAPTIONS.get(key, key)
    png_data = export_chart_as_png(fig)
    if png_data:
        try:
            doc.add_picture(BytesIO(png_data), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(f"图：{caption}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(10)
            cap.runs[0].font.color.rgb = RGBColor(120, 120, 120)
            return None
        except Exception as e:
            return f"（图表「{caption}」嵌入失败：{e}）"
    return f"（图表「{caption}」因缺少 kaleido 包无法嵌入，请在部署环境安装 kaleido）"


# ============================================================
# 封面/标题生成
# ============================================================
def _compute_report_title(has_obs, has_fc):
    """根据数据可用性决定报告标题"""
    if has_obs and has_fc:
        return "气象数据与预报综合分析报告"
    elif has_fc:
        return "数值预报分析报告"
    else:
        return "气象数据分析报告"


def _build_cover_meta(df, fc_df, fc_analysis, warnings_list, score, source):
    """构建封面信息表的所有元数据行"""
    rows = []
    rows.append(("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    if source:
        rows.append(("数据来源", source))
    has_obs = df is not None and not df.empty
    has_fc = fc_df is not None or fc_analysis is not None
    if has_obs:
        ts_col = _resolve_field(df, "timestamp")
        if ts_col and len(df) > 0:
            try:
                rows.append(("观测时段",
                             f"{df[ts_col].min()} ~ {df[ts_col].max()}"))
            except Exception:
                pass
        rows.append(("观测记录数", f"{len(df)} 条"))
        rows.append(("数据质量评分", f"{score}/100"))
        if warnings_list:
            rows.append(("检出预警事件", f"{len(warnings_list)} 个"))
    if has_fc and fc_df is not None and not fc_df.empty:
        fts_col = _resolve_field(fc_df, "timestamp")
        if fts_col and len(fc_df) > 0:
            try:
                rows.append(("预报时段",
                             f"{fc_df[fts_col].min()} ~ {fc_df[fts_col].max()}"))
            except Exception:
                pass
        if "latitude" in fc_df.columns and "longitude" in fc_df.columns:
            try:
                lat = fc_df["latitude"].dropna().iloc[0]
                lon = fc_df["longitude"].dropna().iloc[0]
                if pd.notna(lat) and pd.notna(lon):
                    rows.append(("预报坐标", f"{lat:.4f}°N, {lon:.4f}°E"))
            except Exception:
                pass
    if has_fc and fc_analysis:
        summary = fc_analysis.get("summary", "")
        if summary:
            # 限制封面总述长度
            rows.append(("预报总述", summary[:120] + ("…" if len(summary) > 120 else "")))
    return rows


# ============================================================
# 专业版生成器
# ============================================================
def _build_professional_report(doc, df, fc_df, fc_analysis, life_indices,
                               warnings_list, score, source):
    """专业版：完整章节 + 表格 + 图表 + 防御建议"""
    has_obs = df is not None and not df.empty
    has_fc = fc_df is not None or fc_analysis is not None
    has_warnings = bool(warnings_list)

    # ---- 封面 ----
    title_text = _compute_report_title(has_obs, has_fc)
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _build_cover_table(doc, _build_cover_meta(
        df, fc_df, fc_analysis, warnings_list, score, source))

    # ---- 目录 ----
    toc_sections = [("一、报告说明", "")]
    if has_obs:
        toc_sections.append(("二、观测数据统计摘要", ""))
        if warnings_list:
            toc_sections.append(("三、事件检测与预警", ""))
            toc_sections.append(("七、防御建议", ""))
    if has_fc:
        toc_sections.append(("四、数值预报概况", ""))
    if life_indices:
        toc_sections.append(("五、生活出行指南", ""))
    toc_sections.append(("六、技术说明", ""))
    _add_page_break(doc)
    _build_toc(doc, toc_sections)

    # ---- 一、报告说明 ----
    _add_page_break(doc)
    section_num = 1
    doc.add_heading(f"{_number(section_num)}、报告说明", level=1)
    doc.add_paragraph(
        "本报告基于观测数据和/或数值预报数据自动生成，结合统计、图表、事件检测、"
        "生活指南、防御建议等多个维度，为气象工作者和决策者提供专业的分析参考。"
    )
    doc.add_paragraph(
        "数据处理遵循标准质量控制流程，事件检测采用国家预警分级标准，"
        "生活指数参照气象行业常用分级方法。预报数据使用 NOAA GFS 全球预报系统，"
        "免注册、可商用、空间分辨率约 0.25°，最长预报周期 16 天。"
    )
    doc.add_paragraph(
        "本报告中的图表均嵌入 Word 文档，无需外部依赖即可查看；"
        "如需更高分辨率版本，可在「图表独立下载」区单独导出 PNG。"
    )

    # ---- 二、观测数据统计摘要 ----
    section_num += 1
    if has_obs:
        _add_page_break(doc)
        doc.add_heading(f"{_number(section_num)}、观测数据统计摘要", level=1)

        stats_fields = ["temperature", "pressure", "humidity",
                        "wind_speed", "visibility", "precipitation"]
        labels = {
            "temperature": "气温 (℃)", "pressure": "气压 (hPa)",
            "humidity": "湿度 (%)", "wind_speed": "风速 (m/s)",
            "visibility": "能见度 (km)", "precipitation": "降水量 (mm)",
        }
        resolved = _resolve_stats(df, stats_fields)
        if resolved:
            rows = []
            for f in stats_fields:
                if f in resolved:
                    s = resolved[f]
                    rows.append([
                        labels.get(f, f),
                        f"{s.mean():.2f}",
                        f"{s.min():.2f}",
                        f"{s.max():.2f}",
                        f"{s.std():.2f}",
                        str(len(s)),
                    ])
            _add_table(doc,
                       headers=["字段", "均值", "最小", "最大", "标准差", "记录数"],
                       rows=rows,
                       col_widths=[3, 2.2, 2.2, 2.2, 2.2, 2.2])

        # 观测图表
        obs_chart_figs = {k: v for k, v in _generate_report_charts(df).items()
                          if k.startswith("obs:")}
        if obs_chart_figs:
            doc.add_paragraph("")
            doc.add_heading("观测数据图表", level=2)
            for key, fig in obs_chart_figs.items():
                err = _insert_chart(doc, key, fig)
                if err:
                    doc.add_paragraph(err)

    # ---- 三、事件检测与预警 ----
    if has_obs and warnings_list:
        section_num += 1
        _add_page_break(doc)
        doc.add_heading(f"{_number(section_num)}、事件检测与预警", level=1)
        doc.add_paragraph(
            "以下事件根据国家预警分级标准（《气象灾害预警信号发布与传播办法》）"
            "自动识别，按等级由高到低排列。")
        _render_warnings_table_pro(doc, warnings_list)

    # ---- 四、数值预报概况 ----
    if has_fc:
        section_num += 1
        _add_page_break(doc)
        doc.add_heading(f"{_number(section_num)}、数值预报概况", level=1)
        if fc_analysis:
            _render_forecast_section(doc, fc_analysis)
        # fc 图表
        fc_chart_figs = {k: v for k, v in
                         _generate_report_charts(forecast_df=fc_df).items()
                         if k.startswith("fc:")}
        if fc_chart_figs:
            doc.add_paragraph("")
            doc.add_heading("预报图表", level=2)
            for key, fig in fc_chart_figs.items():
                err = _insert_chart(doc, key, fig)
                if err:
                    doc.add_paragraph(err)

    # ---- 五、生活出行指南 ----
    if life_indices:
        section_num += 1
        _add_page_break(doc)
        doc.add_heading(f"{_number(section_num)}、生活出行指南", level=1)
        doc.add_paragraph(
            "基于预报数据计算的 7 项生活指数，覆盖穿衣、带伞、体感、运动、"
            "紫外线、洗车、晾晒等日常决策场景。")
        _render_life_indices_table(doc, life_indices)

    # ---- 六、技术说明 ----
    section_num += 1
    _add_page_break(doc)
    doc.add_heading(f"{_number(section_num)}、技术说明", level=1)
    doc.add_heading("数据来源", level=2)
    doc.add_paragraph(
        "• 观测数据：用户上传的 CSV/Excel 气象观测记录\n"
        "• 预报数据：Open-Meteo GFS（全球预报系统），空间分辨率 0.25°，"
        "时间分辨率 1 小时，最长 16 天"
    )
    doc.add_heading("预警分级标准", level=2)
    doc.add_paragraph(
        "采用国家预警四级体系：蓝色（IV级，一般）→ 黄色（III级，较重）"
        "→ 橙色（II级，严重）→ 红色（I级，特别严重）。"
        "具体阈值参见 config.py 中 WARN_RULES 定义。"
    )
    doc.add_heading("生活指数计算方法", level=2)
    doc.add_paragraph(
        "穿衣指数：基于气温分 5 级（0-5）\n"
        "带伞建议：基于 24h 降水概率分 3 级（0-3）\n"
        "体感舒适度：Thom 不适指数（SSD = T - 0.55·(1-RH)·(T-14) - V^(1/3)·(T-10)/20）\n"
        "运动指数：基于温湿度组合分 0-100\n"
        "紫外线：基于 UV 指数分 4 级（弱/中/强/很强）\n"
        "洗车指数：基于降水预报分 3 级\n"
        "晾晒指数：基于降水+湿度组合分 0-100"
    )

    # ---- 七、防御建议 ----
    if has_obs and has_warnings:
        section_num += 1
        _add_page_break(doc)
        doc.add_heading(f"{_number(section_num)}、防御建议", level=1)
        _render_defense_advice(doc, warnings_list)

    # ---- 无数据友好提示 ----
    if not has_obs and not has_fc:
        doc.add_paragraph("")
        doc.add_paragraph("⚠ 当前没有数据可生成报告。请先导入观测数据或生成数值预报。")


def _render_warnings_table_pro(doc, warnings_list):
    """专业版预警表格"""
    from config import WARN_LEVEL_ORDER
    sorted_w = sorted(warnings_list,
                      key=lambda w: WARN_LEVEL_ORDER.get(w["level"], 4))
    color_map = {
        "蓝色": RGBColor(0, 102, 204),
        "黄色": RGBColor(245, 166, 35),
        "橙色": RGBColor(242, 101, 34),
        "红色": RGBColor(208, 2, 27),
    }
    rows = []
    for w in sorted_w:
        rows.append([
            f"[{w['level']}] {w['level_num']}",
            w["type"],
            w["detail"][:80] + ("…" if len(w["detail"]) > 80 else ""),
        ])
    _add_table(doc,
               headers=["等级", "类型", "详情"],
               rows=rows,
               col_widths=[3.5, 2.5, 9])
    # 在表格后逐条标注颜色
    for w in sorted_w:
        p = doc.add_paragraph()
        run = p.add_run(f"  → {w['icon']} {w['type']}{w['level']}事件：{w['detail']}")
        run.font.size = Pt(10)
        if w["level"] in color_map:
            run.font.color.rgb = color_map[w["level"]]


def _render_forecast_section(doc, fc_analysis):
    """渲染数值预报概况（不含图表）"""
    summary = fc_analysis.get("summary", "")
    if summary:
        doc.add_paragraph(f"总述：{summary}")

    # 极值表
    ex = fc_analysis.get("extremes", {})
    if ex:
        doc.add_heading("预报期极值", level=2)
        t_max = ex.get("max_temp", (0, ""))
        t_min = ex.get("min_temp", (0, ""))
        max_wind = ex.get("max_wind", (0, ""))
        rows = [
            ["最高气温", f"{t_max[0]:.1f} ℃", str(t_max[1])],
            ["最低气温", f"{t_min[0]:.1f} ℃", str(t_min[1])],
            ["累计降水", f"{ex.get('total_precip', 0):.1f} mm", "—"],
            ["最大风速", f"{max_wind[0]:.1f} m/s", str(max_wind[1])],
        ]
        if "max_daily_precip" in ex:
            mdp = ex["max_daily_precip"]
            rows.append(["最大日降水", f"{mdp[0]:.1f} mm", str(mdp[1])])
        _add_table(doc,
                   headers=["指标", "数值", "日期/说明"],
                   rows=rows,
                   col_widths=[4, 4, 7])

    # 趋势与精度
    prec = fc_analysis.get("precision", {})
    if prec:
        doc.add_heading("趋势与精度详情", level=2)
        rows = []
        tt = prec.get("temp_trend", {})
        if tt:
            rows.append([
                "趋势变化",
                f"{tt.get('diff_mean', 0):+.1f} ℃",
                f"±{tt.get('diff_std', 0):.1f} ℃",
            ])
            rows.append([
                "波动程度",
                tt.get("volatility", ""),
                f"标准差 {tt.get('overall_std', 0):.1f} ℃",
            ])
        rows.append([
            "连续高温",
            f"{prec.get('consecutive_hot', 0)} 天",
            f"最长 {prec.get('consecutive_hot', 0)} 天"
            if prec.get('consecutive_hot', 0) > 0 else "未出现",
        ])
        if "consecutive_rain" in prec:
            rows.append([
                "连续降水",
                f"{prec['consecutive_rain']} 天",
                "累计最长连续降水",
            ])
        diurnal = prec.get("diurnal", {})
        if diurnal:
            rows.append([
                "日较差",
                f"最大 {diurnal.get('max_range', (0, ''))[0]:.1f} ℃",
                f"平均 {diurnal.get('mean_range', 0):.1f} ℃",
            ])
        if rows:
            _add_table(doc,
                       headers=["指标", "数值", "说明"],
                       rows=rows,
                       col_widths=[4, 4, 7])

    # 预警
    fw = fc_analysis.get("warnings", [])
    if fw:
        doc.add_heading("预报预警信号", level=2)
        for w in fw:
            doc.add_paragraph(
                f"- [{w['level']}] {w['type']}: {w['detail']}",
                style="List Bullet")

    # 出行/农业建议
    recs = fc_analysis.get("recommendations", {})
    if recs:
        travel = recs.get("travel", [])
        agri = recs.get("agri", [])
        if travel or agri:
            doc.add_heading("建议", level=2)
        if travel:
            doc.add_paragraph("出行建议：")
            for r in travel[:6]:
                doc.add_paragraph(f"- {r}", style="List Bullet")
        if agri:
            doc.add_paragraph("农业建议：")
            for r in agri[:6]:
                doc.add_paragraph(f"- {r}", style="List Bullet")

    # 耦合风险
    coupling = fc_analysis.get("coupling", [])
    if coupling:
        doc.add_heading("多要素耦合风险", level=2)
        for c in coupling:
            doc.add_paragraph(
                f"- {c['icon']} {c['type']}（{c['severity']}）：{c['detail']}",
                style="List Bullet")


def _render_life_indices_table(doc, life_indices):
    """生活出行指南 - 表格化"""
    from config import LIFE_INDEX_META
    order = ["clothing", "umbrella", "comfort", "exercise",
             "uv", "carwash", "drying"]
    rows = []
    for key in order:
        if key in life_indices:
            info = life_indices[key]
            icon, name = LIFE_INDEX_META.get(key, ("", key))
            score = info.get("score", "")
            score_text = (f"{score}" if isinstance(score, int)
                          else f"{score:.1f}" if isinstance(score, float) else str(score))
            rows.append([
                f"{icon} {name}",
                info.get("level", ""),
                score_text,
                info.get("advice", ""),
            ])
    _add_table(doc,
               headers=["指标", "等级", "评分", "建议"],
               rows=rows,
               col_widths=[3.5, 2.5, 1.5, 7.5])


def _render_defense_advice(doc, warnings_list):
    """防御建议 - 双列表格 + 数据支撑"""
    from config import PUBLIC_ADVICE, AGRI_ADVICE

    doc.add_heading("公众出行", level=2)
    pub_rows = []
    for warn in warnings_list:
        if (warn["type"] in PUBLIC_ADVICE
                and warn["level"] in PUBLIC_ADVICE[warn["type"]]):
            text = PUBLIC_ADVICE[warn["type"]][warn["level"]]
            support = warn.get("detail", "")
            pub_rows.append([
                f"[{warn['level']}] {warn['type']}",
                text,
                support[:60] + ("…" if len(support) > 60 else ""),
            ])
    if pub_rows:
        _add_table(doc,
                   headers=["事件", "建议", "数据支撑"],
                   rows=pub_rows,
                   col_widths=[3.5, 8, 4])
    else:
        doc.add_paragraph("（暂无公众出行建议）")

    doc.add_heading("农业生产", level=2)
    agri_rows = []
    for warn in warnings_list:
        if (warn["type"] in AGRI_ADVICE
                and warn["level"] in AGRI_ADVICE[warn["type"]]):
            text = AGRI_ADVICE[warn["type"]][warn["level"]]
            support = warn.get("detail", "")
            agri_rows.append([
                f"[{warn['level']}] {warn['type']}",
                text,
                support[:60] + ("…" if len(support) > 60 else ""),
            ])
    if agri_rows:
        _add_table(doc,
                   headers=["事件", "建议", "数据支撑"],
                   rows=agri_rows,
                   col_widths=[3.5, 8, 4])
    else:
        doc.add_paragraph("（暂无农业生产建议）")


# ============================================================
# 通俗版生成器（叙述式 + 生活语言 + 问答式）
# ============================================================
_PLAIN_TYPE_DESC = {
    "高温": "天气会很热，可能会让人中暑、心情烦躁，要多喝水、避免长时间在太阳下暴晒。",
    "寒潮": "天气会突然变得很冷，温度下降得很快，注意添衣保暖、老人小孩尽量减少外出。",
    "大风": "风会很大，可能会吹倒东西、广告牌脱落，骑车走路都要小心，避免在广告牌下停留。",
    "暴雨": "雨会下得很大，可能会有积水、低洼处可能内涝，尽量避免外出，如必须外出注意安全。",
    "大雾": "能见度很低，开车要特别小心、要开雾灯、保持车距，走路也要注意车辆。",
    "霜冻": "夜里温度会降到零度以下，会结冰，注意水管防冻、作物防寒。",
    "冰雹": "天会下冰球，会砸坏车玻璃、屋顶、庄稼，尽快到室内躲避，远离窗户。",
    "雷电": "会有打雷闪电，避免在空旷地或树下躲雨，远离金属物品，关闭电器电源。",
    "暴雪": "雪会下得很大，路上会积雪结冰，少开车、必要外出穿防滑鞋。",
    "霾": "空气会比较脏，敏感人群戴 N95 口罩，减少户外运动。",
    "道路结冰": "路面会结冰，走路开车都容易打滑，谨慎慢行。",
    "干旱": "长时间没下雨，水源紧张，注意节约用水，农作物可能需要灌溉。",
    "台风": "热带气旋会带来狂风暴雨，可能造成严重破坏，密切关注、做好防风准备。",
}


def _plain_event_text(warn):
    """生成单个预警的通俗描述（含数据支撑）"""
    base = _PLAIN_TYPE_DESC.get(warn["type"], "")
    if not base:
        base = f"会有{warn['type']}情况，请关注具体影响。"
    support = warn.get("detail", "")
    if support:
        return f"{base}\n    数据说：{support[:80]}"
    return base


def _build_plain_report(doc, df, fc_df, fc_analysis, life_indices,
                        warnings_list, source):
    """通俗版：叙述式 + 生活语言 + 问答式"""
    has_obs = df is not None and not df.empty
    has_fc = fc_df is not None or fc_analysis is not None
    has_warnings = bool(warnings_list)

    # ---- 封面 ----
    title_text = _compute_report_title(has_obs, has_fc) + "（通俗版）"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_rows = [
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    if source:
        cover_rows.append(("数据来源", source))
    if has_obs:
        ts_col = _resolve_field(df, "timestamp")
        if ts_col and len(df) > 0:
            try:
                cover_rows.append(
                    ("数据时段",
                     f"{df[ts_col].min()} 至 {df[ts_col].max()}"))
            except Exception:
                pass
        cover_rows.append(("数据量", f"{len(df)} 条记录"))
    _build_cover_table(doc, cover_rows)

    # ---- 摘要：回答"这天气怎么样" ----
    _add_page_break(doc)
    doc.add_heading("一句话回答：这天气怎么样？", level=1)
    one_liner = _build_one_liner(df, fc_df, fc_analysis)
    if one_liner:
        p = doc.add_paragraph()
        run = p.add_run(one_liner)
        run.bold = True
        run.font.size = Pt(13)
    doc.add_paragraph("")

    # ---- 一、这几天天气会怎么变？ ----
    if has_fc:
        _add_page_break(doc)
        doc.add_heading("一、这几天天气会怎么变？", level=1)
        _render_plain_forecast(doc, fc_df, fc_analysis)

    # ---- 二、要小心什么？ ----
    if has_warnings or (has_fc and fc_analysis and fc_analysis.get("warnings")):
        _add_page_break(doc)
        doc.add_heading("二、要小心什么？", level=1)
        # 优先用观测期检测到的预警，更贴近用户实际
        all_warnings = []
        if has_warnings:
            all_warnings.extend(warnings_list)
        if has_fc and fc_analysis:
            for w in fc_analysis.get("warnings", []):
                all_warnings.append(w)
        _render_plain_warnings(doc, all_warnings)

    # ---- 三、出门穿什么？要不要带伞？ ----
    if life_indices:
        _add_page_break(doc)
        doc.add_heading("三、出门穿什么？要不要带伞？", level=1)
        _render_plain_life_indices(doc, life_indices)

    # ---- 四、给种地朋友的建议 ----
    if has_warnings:
        _add_page_break(doc)
        doc.add_heading("四、给种地朋友的建议", level=1)
        _render_plain_agri_advice(doc, warnings_list)

    # ---- 五、如果你想知道更多 ----
    _add_page_break(doc)
    doc.add_heading("五、如果你想知道更多", level=1)
    doc.add_paragraph(
        "这份报告的「专业版」里有更详细的数据：每小时的温度、湿度、风速、"
        "降水等具体数字，以及气象专业的统计指标和判断标准。如果你是从事气象、"
        "农业、能源、保险等相关工作的朋友，可以切换到专业版查看。"
    )
    doc.add_paragraph(
        "如果对某个数据有疑问，比如「为什么这次报的气温和上次不一样」「这条建议"
        "是怎么算出来的」，可以参考专业版的「技术说明」一节。"
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("💡 提示：")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(_TITLE_COLOR)
    p.add_run("本报告基于 Open-Meteo GFS 数值预报自动生成，"
              "如有重要决策建议，请同时参考国家气象部门权威预报。")


def _build_one_liner(df, fc_df, fc_analysis):
    """生成一句话总览"""
    parts = []
    if fc_analysis:
        summary = fc_analysis.get("summary", "")
        if summary:
            parts.append(summary)
        # 极值
        ex = fc_analysis.get("extremes", {})
        if ex:
            t_max = ex.get("max_temp", (None, ""))
            t_min = ex.get("min_temp", (None, ""))
            if t_max[0] is not None and t_min[0] is not None:
                if t_max[0] - t_min[0] > 15:
                    parts.append(
                        f"未来温差较大（{t_min[0]:.0f}℃ 到 {t_max[0]:.0f}℃），"
                        "要注意衣物调整。")
                elif t_max[0] >= 35:
                    parts.append("未来会很热，注意防暑。")
                elif t_min[0] <= 0:
                    parts.append("未来会比较冷，注意保暖。")
                else:
                    parts.append("未来气温总体适宜。")
        fw = fc_analysis.get("warnings", [])
        if fw:
            types = sorted({w["type"] for w in fw})
            parts.append(f"未来需要关注 {', '.join(types)} 等天气。")
    if df is not None and not df.empty:
        temp_col = _resolve_field(df, "temperature")
        if temp_col:
            t = df[temp_col].dropna()
            if len(t) > 0:
                avg = t.mean()
                parts.append(
                    f"近期历史平均温度约 {avg:.0f}℃。"
                    + (" 偏热。" if avg > 25 else " 适中。" if avg > 10 else " 偏冷。"))
    return " ".join(parts) if parts else "暂无足够数据生成总览。"


def _render_plain_forecast(doc, fc_df, fc_analysis):
    """通俗版 - 预报概况叙述式"""
    if not fc_analysis:
        doc.add_paragraph("暂无预报分析数据。")
        return

    summary = fc_analysis.get("summary", "")
    if summary:
        doc.add_paragraph(f"【总述】{summary}")

    ex = fc_analysis.get("extremes", {})
    if ex:
        t_max = ex.get("max_temp", (None, ""))
        t_min = ex.get("min_temp", (None, ""))
        if t_max[0] is not None:
            feeling = ""
            if t_max[0] >= 35:
                feeling = "（很热）"
            elif t_max[0] >= 30:
                feeling = "（偏热）"
            elif t_max[0] >= 25:
                feeling = "（暖和）"
            doc.add_paragraph(
                f"【温度】最热的一天：{t_max[1]} 达 {t_max[0]:.0f}℃{feeling}；"
                f"最冷的一天：{t_min[1]} 降到 {t_min[0]:.0f}℃。")
        max_wind = ex.get("max_wind", (None, ""))
        if max_wind[0] is not None and max_wind[0] >= 8:
            doc.add_paragraph(
                f"【风】最大风速出现在 {max_wind[1]}，达 {max_wind[0]:.1f} m/s，"
                "户外作业要小心。")
        total_p = ex.get("total_precip", 0)
        if total_p > 50:
            doc.add_paragraph(
                f"【降水】预报期累计降水 {total_p:.0f} mm，雨量较大，"
                "出门一定要带伞，注意防范积水。")
        elif total_p > 10:
            doc.add_paragraph(
                f"【降水】预报期累计降水 {total_p:.0f} mm，"
                "有一定雨量，建议常备雨具。")
        elif total_p > 0:
            doc.add_paragraph(
                f"【降水】预报期累计降水 {total_p:.0f} mm，雨量不大。")
        else:
            doc.add_paragraph("【降水】预报期内基本无有效降水。")

    prec = fc_analysis.get("precision", {})
    if prec:
        hot = prec.get("consecutive_hot", 0)
        if hot >= 3:
            doc.add_paragraph(
                f"【高温持续】连续 {hot} 天日最高气温 ≥ 35℃，"
                "持续高温对身体负担较大，老人小孩尤其注意。")

    coupling = fc_analysis.get("coupling", [])
    if coupling:
        doc.add_paragraph("【特别提示】")
        for c in coupling:
            doc.add_paragraph(f"  • {c['type']}（{c['severity']}）：{c['detail']}",
                              style="List Bullet")


def _render_plain_warnings(doc, all_warnings):
    """通俗版 - 预警信号问答式"""
    if not all_warnings:
        doc.add_paragraph("目前没有触发预警信号，可以放心安排活动。")
        return

    # 按等级排序
    from config import WARN_LEVEL_ORDER
    sorted_w = sorted(all_warnings,
                      key=lambda w: WARN_LEVEL_ORDER.get(w["level"], 4))

    level_intro = {
        "红色": "🔴 红色预警（特别严重）——请立即采取行动，避免不必要外出：",
        "橙色": "🟠 橙色预警（严重）——尽量避免相关活动，做好防护：",
        "黄色": "🟡 黄色预警（较重）——有所准备，减少不必要外出：",
        "蓝色": "🔵 蓝色预警（一般）——关注即可，正常生活工作：",
    }
    for warn in sorted_w:
        intro = level_intro.get(warn["level"], "")
        if intro:
            p = doc.add_paragraph()
            run = p.add_run(intro)
            run.bold = True
        text = _plain_event_text(warn)
        doc.add_paragraph(text)


def _render_plain_life_indices(doc, life_indices):
    """通俗版 - 生活指南叙述化"""
    from config import LIFE_INDEX_META
    order = ["clothing", "umbrella", "comfort", "exercise",
             "uv", "carwash", "drying"]
    guides = {
        "clothing": "出门穿什么？",
        "umbrella": "要不要带伞？",
        "comfort": "体感舒服吗？",
        "exercise": "能不能户外运动？",
        "uv": "紫外线强不强？",
        "carwash": "今天能洗车吗？",
        "drying": "适合晾晒衣物吗？",
    }
    for key in order:
        if key in life_indices:
            info = life_indices[key]
            icon, name = LIFE_INDEX_META.get(key, ("", key))
            question = guides.get(key, name)
            p = doc.add_paragraph()
            run = p.add_run(f"  【{question}】")
            run.bold = True
            doc.add_paragraph(
                f"      {icon} {name}：{info.get('level', '')} — {info.get('advice', '')}")


def _render_plain_agri_advice(doc, warnings_list):
    """通俗版 - 农业建议通俗化"""
    from config import AGRI_ADVICE
    agri_rows = []
    for warn in warnings_list:
        if (warn["type"] in AGRI_ADVICE
                and warn["level"] in AGRI_ADVICE[warn["type"]]):
            agri_rows.append((warn, AGRI_ADVICE[warn["type"]][warn["level"]]))
    if not agri_rows:
        doc.add_paragraph("近期天气对农业生产没有特别不利的影响，可以按计划安排农事。")
        return
    for warn, text in agri_rows:
        p = doc.add_paragraph()
        run = p.add_run(f"  【{warn['type']}{warn['level']}】")
        run.bold = True
        doc.add_paragraph(f"      {text}")
        if warn.get("detail"):
            doc.add_paragraph(
                f"      （数据支撑：{warn['detail'][:60]}）")


# ============================================================
# 主入口
# ============================================================
def export_report_word(df, warnings_list, score, source="",
                       forecast_df=None, forecast_analysis=None,
                       plain_language=False, life_indices=None):
    """生成 Word 分析报告主入口。

    自动根据数据可用性选择章节，字段名自适应。
    plain_language=True 时输出叙述式通俗版，否则输出完整专业版。
    """
    try:
        from docx import Document
    except ImportError:
        st.error("python-docx 未安装，无法导出 Word 报告")
        return None

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    _set_default_styles(doc)

    if plain_language:
        _build_plain_report(doc, df, forecast_df, forecast_analysis,
                            life_indices, warnings_list, source)
    else:
        _build_professional_report(doc, df, forecast_df, forecast_analysis,
                                    life_indices, warnings_list, score, source)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 导出 Tab 渲染
# ============================================================
def render_export_tab(df, warnings_list, score, source=""):
    """渲染增强版报告导出 Tab"""
    st.subheader("[导出] 报告导出")

    # ---- 数据导出 ----
    st.write("#### [统计] 数据导出")
    c1, c2, c3 = st.columns(3)
    with c1:
        if df is not None and not df.empty:
            csv_data = export_data_csv(df)
            if csv_data:
                st.download_button(
                    label="[下载] 处理后数据 (CSV)",
                    data=csv_data,
                    file_name=f"气象数据_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
    with c2:
        fc_df = st.session_state.get("fc_df", None)
        if fc_df is not None and not fc_df.empty:
            fc_csv = export_data_csv(fc_df)
            if fc_csv:
                st.download_button(
                    label="[下载] GFS 预报数据 (CSV)",
                    data=fc_csv,
                    file_name=f"GFS预报_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
    with c3:
        pass

    # ---- Word 报告 ----
    st.write("---")
    st.write("#### [文档] Word 分析报告 (含图表 + 统计 + 预报)")

    report_style = st.radio("报告风格", ["📊 专业版", "💬 通俗版"],
                            horizontal=True, key="report_style")
    if report_style == "💬 通俗版":
        st.caption("用生活化语言叙述数据，减少术语，适合非专业用户阅读。"
                   "包含「这天气怎么样」一句话回答、问答式预警、出门穿衣指南等。")
    else:
        st.caption("完整章节结构：封面 + 目录 + 报告说明 + 观测统计 + 预报概况 "
                   "+ 事件检测 + 生活指南 + 防御建议 + 技术说明。表格化排版，"
                   "图表自动嵌入。")

    fc_analysis = st.session_state.get("fc_analysis", None)

    btn_label = ("[生成] 生成通俗版叙述报告"
                 if report_style == "💬 通俗版"
                 else "[生成] 生成专业版图文报告")
    if st.button(btn_label, use_container_width=True, key="gen_report"):
        with st.spinner("正在生成图文报告..."):
            try:
                doc_data = export_report_word(
                    df, warnings_list, score, source,
                    forecast_df=st.session_state.get("fc_df"),
                    forecast_analysis=fc_analysis,
                    plain_language=(report_style == "💬 通俗版"),
                    life_indices=st.session_state.get("life_indices"),
                )
                if doc_data:
                    st.session_state["report_data"] = doc_data
                    st.success("报告已生成，点击下方按钮下载。")
                    st.rerun()
            except Exception as e:
                st.error(f"报告生成失败: {e}")
                if st.session_state.get("debug_mode"):
                    st.code(traceback.format_exc(), language="python")

    if "report_data" in st.session_state:
        st.download_button(
            label="[下载] 下载 Word 分析报告",
            data=st.session_state["report_data"],
            file_name=f"气象分析报告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    # ---- 图表 PNG 下载 ----
    fc_df = st.session_state.get("fc_df", None)
    has_obs = df is not None and not df.empty
    has_fc = fc_df is not None and not fc_df.empty
    if has_obs or has_fc:
        st.write("---")
        st.write("#### [图表] 独立图表 PNG 下载")
        chart_figs = _generate_report_charts(df=df, forecast_df=fc_df)
        if chart_figs:
            names_display = {
                "obs:temperature": "观测气温",
                "obs:precipitation": "观测降水",
                "obs:pressure": "观测气压",
                "obs:wind_speed": "观测风速",
                "fc:temperature": "GFS 气温/体感",
                "fc:precipitation": "GFS 逐时降水",
                "fc:wind_speed": "GFS 风速",
                "fc:humidity": "GFS 湿度",
                "fc:precip_prob": "GFS 降水概率",
            }
            cols = st.columns(min(len(chart_figs), 5))
            for idx, (name, fig) in enumerate(chart_figs.items()):
                png_bytes = export_chart_as_png(fig)
                display = names_display.get(name, name)
                with cols[idx % len(cols)]:
                    st.caption(display)
                    if png_bytes:
                        st.download_button(
                            label=f"[图片] 下载 PNG",
                            data=png_bytes,
                            file_name=(f"{display}_"
                                       f"{datetime.now().strftime('%Y%m%d')}.png"),
                            mime="image/png",
                            use_container_width=True,
                            key=f"png_{name}",
                        )
                        st.plotly_chart(fig, use_container_width=True,
                                        key=f"rpt_chart_{name}")
                    else:
                        st.info("kaleido 未安装，无法导出图片")

    # ---- 统计摘要 ----
    if df is not None and not df.empty:
        stats_fields = ["temperature", "pressure", "humidity",
                        "wind_speed", "visibility", "precipitation"]
        label_map = {"temperature": "气温", "pressure": "气压", "humidity": "湿度",
                     "wind_speed": "风速", "visibility": "能见度",
                     "precipitation": "降水量"}
        resolved = _resolve_stats(df, stats_fields)
        if resolved:
            st.write("---")
            st.write("#### [列表] 快速统计摘要")
            cols = st.columns(min(len(resolved), 6))
            for i, f in enumerate(stats_fields):
                if f in resolved:
                    with cols[i % len(cols)]:
                        st.metric(label_map.get(f, f), f"{resolved[f].mean():.1f}")