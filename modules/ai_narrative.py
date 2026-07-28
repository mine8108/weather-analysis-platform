"""
AI 预警叙事模块（C 方案）

把 analyzer.py 的国标事件检测、多要素耦合、空气质量评估结果，
转化为面向公众与基层管理人员的自然语言报告。

设计约束（来自方案定稿）：
- 仅生成报告（不含聊天追问），降低复杂度
- 纯 requests 调用，无新增第三方依赖（requests 已在 requirements.txt）
- 密钥走 st.secrets，绝不进代码/前端
- 无密钥 / 调用失败时优雅降级，绝不阻断主流程
"""

import streamlit as st
import requests
import html
import io
import re
from datetime import datetime


# ============================================================
# 一、提示词构造
# ============================================================
def build_prompt(detection):
    """依据检测结果构造发送给 LLM 的提示词。

    detection 结构:
    {
        "warnings":   [ {type, level, level_num, detail, icon}, ... ],   # 8 类国标事件
        "coupling":   [ {type, severity, detail, icon}, ... ],          # 耦合风险
        "air_quality": {aqi, primary, level, color, advice, details} | None,   # 观测污染物
        "air_quality_nwp": {peak_aqi, peak_time, mean_aqi, primary, pollutants, good_hours, bad_hours} | None,
        "wind_nwp":   {dominant, dominant_pct, calm_pct, max_wind_dir, shift} | None,  # 数值预报风况
        "wind_obs":   {dominant, dominant_pct, calm_pct, max_wind_dir, shift} | None,  # 观测风况
        "nwp":        {...} | None,
    }
    返回: 纯文本提示词（含系统指令 + 结构化数据）
    """
    warnings = detection.get("warnings", []) or []
    coupling = detection.get("coupling", []) or []
    aq = detection.get("air_quality")
    nwp = detection.get("nwp")

    sections = []
    nwp_hint = ""
    if nwp:
        nwp_hint = (
            f"下面给出的全部是未来短期（{nwp['period']['hours']} 小时）"
            "数值预报模式（GFS）的客观数据摘要，并非实测，也非气候统计。"
        )
    sections.append(
        "你是一名气象防灾减灾分析助手。下面是一段气象数据，"
        "基于中国国家标准气象预警阈值自动检测的结果。请据此生成一份"
        "面向公众与基层管理人员的自然语言解读报告。"
        + nwp_hint
    )
    sections.append(
        "要求：中文，简洁专业，400-500 字；按【总体概览】【关键预警】"
        "【多要素耦合风险】【风况】【空气质量】【综合建议】六段组织；"
        "若无某项风险，明确说明“未检测到”；"
        "必须严格基于下方给出的具体数值、出现时刻与时段进行解读；"
        "不得编造任何数据；不得将短期预报外推为气候趋势或长期预测；"
        "涉及预警等级仅作参考提示，最终以官方气象部门发布为准；"
        "语气客观、可操作。"
    )

    # 关键预警
    if warnings:
        w_lines = [f"- {w['icon']} {w['type']}{w['level']}（{w['level_num']}）：{w['detail']}" for w in warnings]
        sections.append("【关键预警】\n" + "\n".join(w_lines))
    else:
        sections.append("【关键预警】\n未检测到符合国标阈值的气象预警事件。")

    # 耦合风险
    nwp_coupling_lines = (nwp or {}).get("coupling", []) if nwp else []
    if coupling or nwp_coupling_lines:
        c_lines = [f"- {c['icon']} {c['type']}（{c['severity']}）：{c['detail']}" for c in coupling]
        c_lines += [f"- 数值预报耦合：{c}" for c in nwp_coupling_lines]
        sections.append("【多要素耦合风险】\n" + "\n".join(c_lines))
    else:
        sections.append("【多要素耦合风险】\n未检测到显著的多要素耦合风险。")

    # 风况（合并数值预报与观测风向）
    wind_nwp = detection.get("wind_nwp")
    wind_obs = detection.get("wind_obs")
    wind_lines = []
    if wind_nwp:
        wn = [f"数值预报主导风向：{wind_nwp['dominant']}风（占比 {wind_nwp['dominant_pct']}%）"]
        if wind_nwp.get("max_wind_dir"):
            wn.append(f"最大风速时段风向：{wind_nwp['max_wind_dir']}风")
        if wind_nwp.get("calm_pct"):
            wn.append(f"静风占比 {wind_nwp['calm_pct']}%")
        if wind_nwp.get("shift"):
            wn.append("预报期内风向发生明显转变（前后段主导风向不同）")
        wind_lines.append("【数值预报】" + "；".join(wn) + "。")
    if wind_obs:
        wo = [f"观测主导风向：{wind_obs['dominant']}风（占比 {wind_obs['dominant_pct']}%）"]
        if wind_obs.get("max_wind_dir"):
            wo.append(f"最大风速时段风向：{wind_obs['max_wind_dir']}风")
        if wind_obs.get("calm_pct"):
            wo.append(f"静风占比 {wind_obs['calm_pct']}%")
        if wind_obs.get("shift"):
            wo.append("观测期内风向发生明显转变")
        wind_lines.append("【观测数据】" + "；".join(wo) + "。")
    if wind_lines:
        sections.append("【风况】\n" + "\n".join(wind_lines))
    else:
        sections.append("【风况】\n当前数据未含风向字段，无法评估风况。")

    # 空气质量（合并观测与数值预报）
    aq_nwp = detection.get("air_quality_nwp")
    aq_lines = []
    if aq:
        details = aq.get("details", []) or []
        d_lines = [
            f"- {d['label']}：均值 {d['avg']} μg/m³，IAQI {d['iaqi']}，等级 {d['level']}"
            f"{'（超标）' if (d.get('exceed_daily') or d.get('exceed_hourly')) else ''}"
            for d in details
        ]
        aq_lines.append(
            f"观测数据：综合 AQI {aq['aqi']}，等级 {aq['level']}，首要污染物 {aq.get('primary') or '无'}。\n"
            + "\n".join(d_lines)
        )
    if aq_nwp:
        nwp_p = "；".join(aq_nwp.get("pollutants", [])) if aq_nwp.get("pollutants") else "无逐污染物明细"
        aq_lines.append(
            f"数值预报：峰值 AQI {aq_nwp['peak_aqi']}（{aq_nwp['peak_time']}），"
            f"均值 {aq_nwp['mean_aqi']:.0f}，首要污染物 {aq_nwp['primary']}；"
            f"优良时段 {aq_nwp['good_hours']} 小时、污染时段（AQI>100）{aq_nwp['bad_hours']} 小时。"
            f"逐污染物：{nwp_p}。"
        )
    if aq_lines:
        sections.append("【空气质量】\n" + "\n".join(aq_lines))
    else:
        sections.append("【空气质量】\n当前数据未含大气污染物字段，无法评估空气质量。")

    # 数值预报结构化事实（严格依据，禁止外推）
    if nwp:
        p = nwp.get("period", {})
        fact_lines = [
            f"预报时效：{p.get('start')} 至 {p.get('end')}，共 {p.get('hours')} 小时。",
        ]
        t = nwp.get("temperature")
        if t:
            fact_lines.append(
                f"气温：最高 {t['max']}℃（{t['max_time']}），最低 {t['min']}℃（{t['min_time']}），"
                f"平均 {t['mean']}℃；≥35℃ {t['hot_hours']} 小时，≥37℃ {t['severe_hot_hours']} 小时。"
            )
        at = nwp.get("apparent_temperature")
        if at:
            fact_lines.append(f"体感温度最高 {at['max']}℃（{at['max_time']}）。")
        h = nwp.get("humidity")
        if h:
            fact_lines.append(f"相对湿度：{h['min']}%~{h['max']}%（均值 {h['mean']}%）。")
        pr = nwp.get("precipitation")
        if pr:
            fact_lines.append(
                f"降水：累计 {pr['total']}mm，最大 1 小时 {pr['max_1h']}mm（{pr['max_1h_time']}），"
                f"降雨时段 {pr['rain_hours']} 小时，其中大雨量级（>10mm/h）{pr['heavy_hours']} 小时。"
            )
        w = nwp.get("wind_speed")
        if w:
            fact_lines.append(
                f"风速：最大 {w['max']}m/s（{w['max_time']}），"
                f"≥6 级（10.8m/s）{w['gale_hours']} 小时，≥5 级（8m/s）{w['strong_hours']} 小时。"
            )
        nwp_wind = nwp.get("wind")
        if nwp_wind:
            wind_fact = f"风向：主导 {nwp_wind['dominant']}风（占比 {nwp_wind['dominant_pct']}%）"
            if nwp_wind.get("max_wind_dir"):
                wind_fact += f"，最大风速时段 {nwp_wind['max_wind_dir']}风"
            if nwp_wind.get("calm_pct"):
                wind_fact += f"，静风占比 {nwp_wind['calm_pct']}%"
            if nwp_wind.get("shift"):
                wind_fact += "，预报期内风向明显转变"
            fact_lines.append(wind_fact + "。")
        segs = nwp.get("segments", [])
        if segs:
            seg_text = "；".join(
                f"{s['window']} 段：T∈[{s.get('min_temp','-')},{s.get('max_temp','-')}]℃"
                f"、降水 {s.get('total_precip', 0)}mm、最大风 {s.get('max_wind', 0)}m/s"
                for s in segs
            )
            fact_lines.append("分时段： " + seg_text + "。")
        nwp_coupling = nwp.get("coupling", [])
        if nwp_coupling:
            fact_lines.append("多要素耦合： " + "；".join(nwp_coupling) + "。")
        if aq_nwp:
            nwp_p = "；".join(aq_nwp.get("pollutants", [])) or "无"
            fact_lines.append(
                f"空气质量预报：峰值 AQI {aq_nwp['peak_aqi']}（{aq_nwp['peak_time']}），"
                f"均值 {aq_nwp['mean_aqi']:.0f}，首要污染物 {aq_nwp['primary']}；"
                f"优良 {aq_nwp['good_hours']}h / 污染(AQI>100) {aq_nwp['bad_hours']}h；"
                f"逐污染物：{nwp_p}。"
            )
        alerts = nwp.get("alerts", [])
        if alerts:
            al_text = "；".join(f"{a['type']}{a['level']}（依据：{a['basis']}）" for a in alerts)
            fact_lines.append("国标等级初步判定（仅供参考）： " + al_text + "。")
        sections.append("【数值预报事实数据】\n" + "\n".join(fact_lines))

    sections.append("请直接输出报告正文（不要重复系统指令，也不要重复上述事实数据原文，应转化为连贯的自然语言解读）。")
    return "\n\n".join(sections)


# ============================================================
# 二、LLM 调用（DeepSeek，OpenAI 兼容协议）
# ============================================================
def call_llm(prompt, api_key, base_url=None, model=None):
    """调用 LLM 生成解读文本。失败时抛异常，由调用方降级处理。

    密钥来源：st.secrets（LLM_API_KEY / LLM_BASE_URL / LLM_MODEL）。
    """
    base_url = (base_url or st.secrets.get("LLM_BASE_URL", "https://api.deepseek.com")).rstrip("/")
    model = model or st.secrets.get("LLM_MODEL", "deepseek-chat")

    system_msg = (
        "你是严谨的气象分析助手，仅基于用户提供的数据生成解读，"
        "不做无根据的推测，中文输出。"
    )
    resp = requests.post(
        f"{base_url}/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.3,
            "max_tokens": 800,
        },
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"].strip()


# ============================================================
# 三、降级：结构化 Markdown 摘要（无密钥或调用失败时兜底）
# ============================================================
def build_fallback_markdown(detection):
    """无 LLM 时，把结构化结果直接渲染成可读摘要（【段】结构，复用美化渲染）。"""
    warnings = detection.get("warnings", []) or []
    coupling = detection.get("coupling", []) or []
    aq = detection.get("air_quality")
    aq_nwp = detection.get("air_quality_nwp")

    parts = ["【总体概览】\nAI 解读模型不可用，已降级为系统自动生成的要点摘要。"]
    if warnings:
        lines = [f"- {w['icon']} {w['type']}{w['level']}：{w['detail']}" for w in warnings]
        parts.append("【关键预警】\n" + "\n".join(lines))
    else:
        parts.append("【关键预警】\n未检测到符合阈值的事件。")
    if coupling:
        lines = [f"- {c['icon']} {c['type']}（{c['severity']}）：{c['detail']}" for c in coupling]
        parts.append("【多要素耦合风险】\n" + "\n".join(lines))
    aq_lines = []
    if aq:
        aq_lines.append(f"观测：AQI {aq['aqi']}（{aq['level']}），首要污染物 {aq.get('primary') or '无'}。")
    if aq_nwp:
        aq_lines.append(
            f"数值预报：峰值 AQI {aq_nwp['peak_aqi']}（{aq_nwp['peak_time']}），"
            f"均值 {aq_nwp['mean_aqi']:.0f}，首要污染物 {aq_nwp['primary']}。"
        )
    if aq_lines:
        parts.append("【空气质量】\n" + "\n".join(aq_lines))
    return "\n\n".join(parts)


# ============================================================
# 三·五、报告排版与导出（学术白底，网页/文件 1:1）
# ============================================================
def _parse_sections(text):
    """按 【标题】 切片为 [(title, body), ...]；无标记则整体作为「解读摘要」。"""
    parts = re.split(r"【([^】]{1,24})】", text)
    sections = []
    preamble = parts[0].strip()
    if preamble:
        sections.append(("解读摘要", preamble))
    for i in range(1, len(parts), 2):
        title = parts[i].strip()
        body = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if title:
            sections.append((title, body))
    return sections


def _build_meta(detection):
    """推导数据范围与生成时间。"""
    has_nwp = detection.get("nwp") is not None
    has_obs = (
        detection.get("air_quality") is not None
        or detection.get("wind_obs") is not None
        or any(w.get("type") != "数值预报" for w in detection.get("warnings", []))
    )
    if has_nwp and has_obs:
        scope = "观测数据 + 数值预报（GFS）"
    elif has_nwp:
        scope = "数值预报（GFS）"
    else:
        scope = "观测数据"
    return {"scope": scope, "generated": datetime.now().strftime("%Y-%m-%d %H:%M")}


_CSS = """
.report-card{background:#ffffff;color:#1f2933;border:1px solid #e2e8f0;
  border-radius:10px;padding:20px 24px;margin:10px 0;
  box-shadow:0 1px 3px rgba(15,23,42,.10);
  font-family:"Microsoft YaHei","PingFang SC","Source Han Sans SC",sans-serif;
  line-height:1.75;}
.report-head{border-bottom:2px solid #1f4e79;padding-bottom:8px;margin-bottom:14px;}
.report-title{font-size:20px;font-weight:700;color:#1f4e79;letter-spacing:1px;}
.report-sub{font-size:12px;color:#64748b;margin-top:4px;}
.sec{margin:14px 0;}
.sec-h{font-size:15px;font-weight:700;color:#1f4e79;
  border-left:4px solid #1f4e79;padding-left:10px;margin-bottom:6px;}
.report-card p{margin:4px 0;font-size:14px;color:#27303a;}
"""


def _report_html(sections, meta):
    """生成学术风报告 HTML（白底深蓝标题，强制浅色，独立于 app 主题）。"""
    body_html = ""
    for title, body in sections:
        paras = [p for p in body.split("\n") if p.strip()]
        if not paras:
            continue
        p_html = "".join(f"<p>{html.escape(p)}</p>" for p in paras)
        body_html += (
            f'<div class="sec"><div class="sec-h">{html.escape(title)}</div>{p_html}</div>'
        )
    return (
        f'<div class="report-card">'
        f'<div class="report-head">'
        f'<div class="report-title">气象智能解读报告</div>'
        f'<div class="report-sub">生成时间：{html.escape(meta["generated"])}'
        f'　｜　数据范围：{html.escape(meta["scope"])}</div>'
        f"</div>{body_html}</div>"
        f"<style>{_CSS}</style>"
    )


def _set_cjk(run, font="Microsoft YaHei"):
    """为 run 设置中日韩字体（解决 docx 中文宋体回退问题）。"""
    from docx.oxml.ns import qn
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def _build_docx(sections, meta):
    """同内容导出为 .docx（学术白底版式）。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Microsoft YaHei"
    from docx.oxml.ns import qn
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("气象智能解读报告", level=0)
    for r in title.runs:
        _set_cjk(r)
    sub = doc.add_paragraph()
    srun = sub.add_run(f"生成时间：{meta['generated']}    数据范围：{meta['scope']}")
    srun.font.size = Pt(9)
    srun.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    _set_cjk(srun)

    for title_txt, body in sections:
        h = doc.add_heading(title_txt, level=1)
        for r in h.runs:
            _set_cjk(r)
        for p in body.split("\n"):
            if p.strip():
                para = doc.add_paragraph(p.strip())
                for r in para.runs:
                    _set_cjk(r)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 四、UI 渲染入口
# ============================================================
def _display_report(text, detection):
    """统一渲染：美化 HTML 卡片 + .docx 下载按钮。"""
    meta = st.session_state.get("ai_narrative_meta") or _build_meta(detection)
    sections = _parse_sections(text)
    st.markdown(_report_html(sections, meta), unsafe_allow_html=True)

    try:
        docx_bytes = _build_docx(sections, meta)
        st.download_button(
            label="⬇ 下载报告 (.docx)",
            data=docx_bytes,
            file_name="气象智能解读报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="ai_download_docx",
        )
    except Exception as e:  # noqa: BLE001 - 导出失败不影响展示
        st.warning(f"报告生成成功，但 .docx 导出失败：{e}")


def render_ai_block():
    """检测 Tab 末尾的 AI 解读块。依赖 session_state['detection_result']。"""
    detection = st.session_state.get("detection_result")
    if not detection:
        return

    st.write("---")
    st.write("### [AI] 智能预警解读")

    api_key = st.secrets.get("LLM_API_KEY", "")
    if not api_key:
        st.info(
            "AI 解读未启用：请在 Streamlit Cloud 的 Secrets 中配置 "
            "`LLM_API_KEY`（DeepSeek）。配置后将出现「生成 AI 解读报告」按钮。"
        )
        return

    if st.button("生成 AI 解读报告", key="ai_generate_report"):
        with st.spinner("AI 正在生成解读..."):
            try:
                prompt = build_prompt(detection)
                text = call_llm(prompt, api_key)
                st.session_state["ai_narrative_text"] = text
                st.session_state["ai_narrative_meta"] = _build_meta(detection)
            except Exception as e:  # noqa: BLE001 - 任何失败都降级，不阻断
                st.error(f"AI 生成失败，已降级为结构化摘要：{e}")
                text = build_fallback_markdown(detection)
                st.session_state["ai_narrative_text"] = text
                st.session_state["ai_narrative_meta"] = _build_meta(detection)

    cached = st.session_state.get("ai_narrative_text")
    if cached:
        _display_report(cached, detection)
